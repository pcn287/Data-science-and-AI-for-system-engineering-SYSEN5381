from __future__ import annotations

import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


HOMEWORK_DIR = Path(__file__).resolve().parent
OUTPUT_DIR = HOMEWORK_DIR / "outputs"
DOCX_PATH = HOMEWORK_DIR / "homework3_submission.docx"


def add_heading(document: Document, text: str, level: int = 1) -> None:
    document.add_heading(text, level=level)


def add_table_from_dataframe(document: Document, df: pd.DataFrame) -> None:
    table = document.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for idx, column in enumerate(df.columns):
        hdr_cells[idx].text = str(column)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = "" if pd.isna(value) else str(value)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_numbered(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Number")


def main() -> None:
    anova = json.loads((OUTPUT_DIR / "anova_results.json").read_text(encoding="utf-8"))
    criteria = json.loads((HOMEWORK_DIR / "validation_criteria.json").read_text(encoding="utf-8"))
    summary_df = pd.read_csv(OUTPUT_DIR / "prompt_score_summary.csv").round(3)
    validation_df = pd.read_csv(OUTPUT_DIR / "validation_scores.csv")
    stats_md = (OUTPUT_DIR / "descriptive_statistics.md").read_text(encoding="utf-8")

    document = Document()
    styles = document.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Homework 3: AI Report Validation System")
    run.bold = True
    run.font.size = Pt(18)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Data Science and AI for Systems Engineering: APIs, Apps, and Analytics for Cloud Systems")

    document.add_paragraph("Name: Puchun Niu")
    document.add_paragraph("Email: pn287@cornell.edu")
    document.add_paragraph("Submission: Homework 3")

    add_heading(document, "Writing Component", 1)
    document.add_paragraph(
        "For this homework, I built a validation system for AI-generated methane emissions "
        "reports. My use case is GreenFeed visit-level CH4 data from feeder 453. Instead of "
        "asking an AI agent to read a large raw CSV, I first use Python to load the locally "
        "saved emissions file, keep the four fields needed for this analysis, and compute "
        "descriptive statistics. This made the workflow much faster and also made the report "
        "generation more controlled, because the AI receives a compact statistical summary "
        "rather than the full raw dataset."
    )
    document.add_paragraph(
        "The validation system is customized for methane report quality. I did not reuse the "
        "generic Module 9 Likert categories. Instead, I created a 0-10 rubric with four "
        "dimensions: statistical accuracy, emissions interpretation, decision support value, "
        "and transparency. These categories fit my report because I wanted to know whether a "
        "generated report used the methane statistics correctly, explained animal-level "
        "variation clearly, gave useful monitoring recommendations, and stated the evidence "
        "and limitations transparently."
    )
    document.add_paragraph(
        "I compared three prompt variants. Prompt A asked for a basic concise summary, Prompt "
        "B emphasized farm-management decision support, and Prompt C asked for deeper "
        "statistical interpretation. For each prompt, the system generated repeated reports "
        "and then used an AI reviewer to score each report against the custom rubric. I used "
        "the composite score, which averages the four validation dimensions, as the main "
        "outcome for comparing prompt quality."
    )
    document.add_paragraph(
        f"In the current experiment, Prompt C had the highest mean composite score "
        f"({summary_df.loc[summary_df['prompt_variant'] == 'C_statistical_depth', 'mean'].iloc[0]}), "
        f"followed by Prompts A and B. I ran a one-way ANOVA to compare composite scores "
        f"across the prompt variants. The ANOVA result was F = {anova['f_statistic']} with "
        f"p = {anova['p_value']}. This means Prompt C performed best by average score, but "
        "the difference was not statistically significant at alpha = 0.05. The main "
        "challenge was that the AI reviewer scored most reports highly, so the scores had "
        "a narrow range. In a final larger experiment, I would increase the number of runs "
        "per prompt and make the validation rubric even stricter to produce more separation."
    )

    add_heading(document, "Git Repository Links", 1)
    add_bullets(
        document,
        [
            "Validation system script: https://github.com/pcn287/Data-science-and-AI-for-system-engineering-SYSEN5381/blob/main/homework3/Multi-Agent.py",
            "Validation criteria/rubric: https://github.com/pcn287/Data-science-and-AI-for-system-engineering-SYSEN5381/blob/main/homework3/validation_criteria.json",
            "Validation scores: https://github.com/pcn287/Data-science-and-AI-for-system-engineering-SYSEN5381/blob/main/homework3/outputs/validation_scores.csv",
            "ANOVA results: https://github.com/pcn287/Data-science-and-AI-for-system-engineering-SYSEN5381/blob/main/homework3/outputs/anova_results.json",
            "Generated reports: https://github.com/pcn287/Data-science-and-AI-for-system-engineering-SYSEN5381/blob/main/homework3/outputs/generated_reports.csv",
            "Prompt comparison plot: https://github.com/pcn287/Data-science-and-AI-for-system-engineering-SYSEN5381/blob/main/homework3/outputs/prompt_comparison_boxplot.png",
        ],
    )

    add_heading(document, "Screenshots and Output Samples", 1)
    document.add_paragraph(
        "The following output samples show the validation system in action. The script loads "
        "local data, computes descriptive statistics, generates reports from three prompt "
        "variants, validates each report, and runs ANOVA on the composite scores."
    )

    add_heading(document, "Sample Descriptive Statistics", 2)
    for line in stats_md.splitlines()[:18]:
        if line.strip():
            document.add_paragraph(line)

    add_heading(document, "Validation Score Summary", 2)
    add_table_from_dataframe(document, summary_df)

    add_heading(document, "ANOVA Output", 2)
    add_table_from_dataframe(
        document,
        pd.DataFrame(
            [
                {
                    "test": anova["test"],
                    "F statistic": anova["f_statistic"],
                    "p-value": anova["p_value"],
                    "best prompt": anova["best_prompt_by_mean_score"],
                    "significant at 0.05": anova["significant_at_0.05"],
                }
            ]
        ),
    )
    document.add_paragraph(anova["interpretation"])

    boxplot_path = OUTPUT_DIR / "prompt_comparison_boxplot.png"
    if boxplot_path.is_file():
        add_heading(document, "Prompt Comparison Boxplot", 2)
        document.add_picture(str(boxplot_path), width=Inches(6.2))

    add_heading(document, "Validation Criteria Table", 1)
    criteria_rows = []
    for name, info in criteria["criteria"].items():
        criteria_rows.append(
            {
                "Dimension": name,
                "Description": info["description"],
                "Scale": info["scale"],
                "Benchmark": info["benchmark"],
            }
        )
    add_table_from_dataframe(document, pd.DataFrame(criteria_rows))
    document.add_paragraph(
        "These criteria differ from the lab's generic Likert scores because they are tailored "
        "to methane emissions reports and use a 0-10 scale. The rubric checks whether the "
        "report is statistically accurate, interprets CH4 emissions meaningfully, supports "
        "management decisions, and explains its scope and limitations."
    )

    add_heading(document, "Experimental Design", 1)
    document.add_paragraph(
        "The experiment compared three prompt variants for generating methane report "
        "interpretations from the same descriptive statistics."
    )
    add_numbered(
        document,
        [
            "Prompt A_basic_summary: concise homework-style summary.",
            "Prompt B_decision_support: farm-manager decision support focus.",
            "Prompt C_statistical_depth: more detailed statistical interpretation.",
        ],
    )
    document.add_paragraph(
        "The current run collected two validation scores per prompt variant, for six total "
        "validated reports. The script can be rerun with a larger sample size using "
        "`py homework3/Multi-Agent.py --iterations 5`."
    )

    add_heading(document, "Validation Results", 1)
    add_table_from_dataframe(
        document,
        validation_df[
            [
                "prompt_variant",
                "run_id",
                "statistical_accuracy",
                "emissions_interpretation",
                "decision_support_value",
                "transparency",
                "composite_score",
            ]
        ],
    )

    add_heading(document, "Statistical Analysis", 1)
    document.add_paragraph(
        "Hypothesis: at least one prompt variant produces reports with a different mean "
        "validation score. I used one-way ANOVA because there were three prompt groups and "
        "one numeric outcome, the composite validation score."
    )
    document.add_paragraph(
        f"The ANOVA produced F = {anova['f_statistic']} and p = {anova['p_value']}. "
        f"The best prompt by mean score was {anova['best_prompt_by_mean_score']}. "
        "Because the p-value was greater than 0.05, I do not reject the null hypothesis "
        "for this small experiment. The practical interpretation is that the statistically "
        "detailed prompt looked strongest in the sample, but the evidence is not strong "
        "enough to claim a significant difference."
    )

    add_heading(document, "System Design and Technical Details", 1)
    add_bullets(
        document,
        [
            "Main script: homework3/Multi-Agent.py",
            "Local data file: homework3/emissions_visits_fid453_20260204_162120.csv",
            "Rubric file: homework3/validation_criteria.json",
            "Output folder: homework3/outputs/",
            "Required packages: pandas, openai, python-dotenv, scipy, matplotlib",
            "Required environment variable: OPENAI_API_KEY",
        ],
    )

    add_heading(document, "Usage Instructions", 1)
    document.add_paragraph("Install dependencies:")
    document.add_paragraph("py -m pip install -r homework3/requirements.txt")
    document.add_paragraph("Run a quick test:")
    document.add_paragraph("py homework3/Multi-Agent.py --iterations 1")
    document.add_paragraph("Run a larger experiment:")
    document.add_paragraph("py homework3/Multi-Agent.py --iterations 5")

    document.add_page_break()
    add_heading(document, "Appendix: Notes for Final Submission", 1)
    document.add_paragraph(
        "Before submitting, I should confirm the GitHub links work after pushing the homework3 "
        "folder. I can also rerun the experiment with more iterations if I want stronger "
        "statistical evidence and more screenshots."
    )

    document.save(DOCX_PATH)
    print(f"Saved {DOCX_PATH}")


if __name__ == "__main__":
    main()
